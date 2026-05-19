#!/usr/bin/env python3
"""
find_suppliers_for_services.py

Loops over solicitations categorized as 'services' and uses
find_relevant_suppliers.find_vendors_for_notice to discover local suppliers
in each solicitation's place of performance. Writes a Word doc listing up to
--max-results solicitations that have at least one supplier match. Optional
JSON/CSV sidecars for downstream tooling.

Usage:
    python find_suppliers_for_services.py [options]

Options:
    --max-results N     Stop after collecting N matches with vendors (default: 10)
    --state XX          Filter by place-of-performance state (e.g. CA)
    --top-n N           Vendors to return per solicitation (default: 3)
    --max-google N      Google results per query (default: 10)
    --scan-cap N        Max solicitations to scan when searching for matches (default: 200)
    --output FILE       Word doc output (default: services_suppliers.docx)
    --json FILE         Also write JSON (no default)
    --csv FILE          Also write a flat CSV (no default)
    --dry-run           List candidate solicitations and exit

Environment variables (required):
    SUPABASE_DB_URL     PostgreSQL connection string
    GOOGLE_API_KEY      Google Custom Search API key
    GOOGLE_CX           Google Custom Search Engine ID
    OPENAI_API_KEY      OpenAI API key
"""

from __future__ import annotations

import argparse
import csv
import json
import logging
import os
import sys
import time
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt, RGBColor
from sqlalchemy import create_engine, text

try:
    from dotenv import load_dotenv
    env_file = Path(__file__).parent / ".env"
    if env_file.exists():
        load_dotenv(env_file)
except ImportError:
    pass

from find_relevant_suppliers import find_vendors_for_notice

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    datefmt="%H:%M:%S",
)
log = logging.getLogger(__name__)


def load_services_solicitations(db_url: str, state: str | None, scan_cap: int) -> list[dict[str, Any]]:
    engine = create_engine(db_url, pool_pre_ping=True)
    sql = """
        SELECT notice_id, title, description, naics_code,
               pop_city, pop_state, response_date, link
        FROM solicitationraw
        WHERE category = 'services'
    """
    params: dict[str, Any] = {}
    if state:
        sql += " AND UPPER(pop_state) = :state"
        params["state"] = state.upper()
    sql += " ORDER BY response_date NULLS LAST LIMIT :scan_cap"
    params["scan_cap"] = scan_cap

    with engine.connect() as conn:
        rows = conn.execute(text(sql), params).mappings().all()
    return [dict(r) for r in rows]


def write_docx(results: list[dict[str, Any]], out_path: Path, generated_at: str) -> None:
    doc = Document()

    title = doc.add_heading("Services Solicitations — Local Supplier Matches", level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    sub = doc.add_paragraph()
    sub.add_run(f"Generated {generated_at}  •  {len(results)} solicitation(s)").italic = True

    for idx, sol in enumerate(results, 1):
        doc.add_heading(f"{idx}. {sol.get('title') or '(no title)'}", level=1)

        meta = doc.add_paragraph()
        meta.add_run("Notice ID: ").bold = True
        meta.add_run(str(sol.get("notice_id") or "—"))
        meta.add_run("    NAICS: ").bold = True
        meta.add_run(str(sol.get("naics_code") or "—"))
        meta.add_run("    Response due: ").bold = True
        meta.add_run(str(sol.get("response_date") or "—"))

        loc = doc.add_paragraph()
        loc.add_run("Place of performance: ").bold = True
        city = (sol.get("pop_city") or "").strip()
        state = (sol.get("pop_state") or "").strip()
        loc.add_run(", ".join(p for p in [city, state] if p) or "—")

        if sol.get("link"):
            link_p = doc.add_paragraph()
            link_p.add_run("Link: ").bold = True
            link_p.add_run(str(sol["link"]))

        doc.add_paragraph("Recommended local suppliers:").runs[0].bold = True
        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        for cell, label in zip(hdr, ["#", "Supplier", "Website / Location", "Why"]):
            cell.text = ""
            run = cell.paragraphs[0].add_run(label)
            run.bold = True
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        for rank, vendor in enumerate(sol.get("vendors", []), 1):
            row = table.add_row().cells
            row[0].text = str(rank)
            row[1].text = vendor.get("name") or "—"
            web = vendor.get("website") or ""
            vloc = vendor.get("location") or ""
            row[2].text = "\n".join(p for p in [web, vloc] if p) or "—"
            row[3].text = vendor.get("reason") or ""

        # readable font in tables
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(10)

        doc.add_paragraph()  # spacer

    doc.save(out_path)


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    parser.add_argument("--max-results", type=int, default=10)
    parser.add_argument("--state", type=str, default=None)
    parser.add_argument("--top-n", type=int, default=3)
    parser.add_argument("--max-google", type=int, default=10)
    parser.add_argument("--scan-cap", type=int, default=200)
    parser.add_argument("--output", type=str, default="services_suppliers.docx")
    parser.add_argument("--json", type=str, default=None)
    parser.add_argument("--csv", type=str, default=None)
    parser.add_argument("--dry-run", action="store_true")
    args = parser.parse_args()

    db_url = os.getenv("SUPABASE_DB_URL")
    google_key = os.getenv("GOOGLE_API_KEY")
    google_cx = os.getenv("GOOGLE_CX")
    openai_key = os.getenv("OPENAI_API_KEY")

    missing = [n for n, v in [
        ("SUPABASE_DB_URL", db_url),
        ("GOOGLE_API_KEY", google_key),
        ("GOOGLE_CX", google_cx),
        ("OPENAI_API_KEY", openai_key),
    ] if not v]
    if missing and not args.dry_run:
        log.error(f"Missing required env vars: {', '.join(missing)}")
        return 1
    if not db_url:
        log.error("SUPABASE_DB_URL required")
        return 1

    sols = load_services_solicitations(db_url, args.state, args.scan_cap)
    log.info(f"Scanning {len(sols)} services solicitation(s); target {args.max_results} matches with vendors")

    if args.dry_run:
        for s in sols:
            log.info(f"  {s['notice_id']} | {s.get('pop_state') or '--'} | {(s.get('title') or '')[:80]}")
        return 0

    results: list[dict[str, Any]] = []
    skipped_no_vendors = 0

    for i, sol in enumerate(sols, 1):
        if len(results) >= args.max_results:
            log.info(f"Reached --max-results={args.max_results}; stopping scan")
            break

        title = (sol.get("title") or "")[:80]
        log.info(f"[scan {i}/{len(sols)} | matched {len(results)}/{args.max_results}] {sol['notice_id']} — {title}")
        try:
            df, _ = find_vendors_for_notice(
                sol,
                google_api_key=google_key,
                google_cx=google_cx,
                openai_api_key=openai_key,
                max_google=args.max_google,
                top_n=args.top_n,
            )
            vendors = df.to_dict(orient="records") if not df.empty else []
        except Exception as exc:
            log.warning(f"  vendor lookup failed: {exc}")
            vendors = []

        if not vendors:
            log.info("  → no suppliers found, skipping")
            skipped_no_vendors += 1
            time.sleep(0.3)
            continue

        results.append({
            "notice_id": sol["notice_id"],
            "title": sol.get("title"),
            "naics_code": sol.get("naics_code"),
            "pop_city": sol.get("pop_city"),
            "pop_state": sol.get("pop_state"),
            "response_date": sol.get("response_date"),
            "link": sol.get("link"),
            "vendors": vendors,
        })
        log.info(f"  → {len(vendors)} vendor(s) [match {len(results)}/{args.max_results}]")
        time.sleep(0.5)

    log.info(f"Done. Matched {len(results)}; skipped {skipped_no_vendors} for lack of suppliers")

    if not results:
        log.warning("No solicitations had supplier matches — no document written")
        return 0

    generated_at = datetime.now().strftime("%Y-%m-%d %H:%M")
    out_path = Path(args.output)
    write_docx(results, out_path, generated_at)
    log.info(f"Wrote {out_path}")

    if args.json:
        Path(args.json).write_text(json.dumps(results, indent=2, default=str))
        log.info(f"Wrote {args.json}")

    if args.csv:
        csv_rows = []
        for sol in results:
            for v in sol["vendors"]:
                csv_rows.append({
                    "notice_id": sol["notice_id"],
                    "solicitation_title": sol.get("title"),
                    "pop_city": sol.get("pop_city"),
                    "pop_state": sol.get("pop_state"),
                    "vendor_name": v.get("name"),
                    "vendor_website": v.get("website"),
                    "vendor_location": v.get("location"),
                    "reason": v.get("reason"),
                })
        with open(args.csv, "w", newline="") as f:
            writer = csv.DictWriter(f, fieldnames=list(csv_rows[0].keys()))
            writer.writeheader()
            writer.writerows(csv_rows)
        log.info(f"Wrote {args.csv} ({len(csv_rows)} rows)")

    return 0


if __name__ == "__main__":
    sys.exit(main())
