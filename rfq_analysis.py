#!/usr/bin/env python3
"""
RFQ Analysis Script
-------------------
Reads RFQ line items from the database, uses AI to categorize parts,
and generates a summary document showing frequently appearing part types.

Usage:
    python rfq_analysis.py
    python rfq_analysis.py --output custom_report.docx
"""

import os
import sys
import json
from datetime import datetime
from dotenv import load_dotenv
from sqlalchemy import create_engine, text
from openai import OpenAI

load_dotenv()

DATABASE_URL = os.environ.get("SUPABASE_DB_URL")
OPENAI_API_KEY = os.environ.get("OPENAI_API_KEY")


def get_db_engine():
    if not DATABASE_URL:
        raise ValueError("Missing SUPABASE_DB_URL environment variable.")
    return create_engine(DATABASE_URL)


def fetch_all_line_items(engine):
    """Fetch all line items from the database."""
    with engine.connect() as conn:
        result = conn.execute(text("""
            SELECT 
                li.part_number,
                li.description,
                li.quantity,
                li.naics_code,
                r.rfq_number,
                r.due_date
            FROM rfq_line_items li
            JOIN rfqs r ON li.rfq_id = r.id
            ORDER BY r.due_date DESC
        """))
        return [dict(row._mapping) for row in result]


def categorize_parts_with_ai(line_items: list) -> dict:
    """Use OpenAI to categorize parts into logical groups."""
    if not OPENAI_API_KEY:
        raise ValueError("Missing OPENAI_API_KEY environment variable.")

    client = OpenAI(api_key=OPENAI_API_KEY)

    # Build a list of unique parts with their descriptions
    unique_parts = {}
    for item in line_items:
        pn = item["part_number"]
        if pn not in unique_parts:
            unique_parts[pn] = {
                "part_number": pn,
                "description": item["description"] or "",
                "count": 0,
                "total_quantity": 0
            }
        unique_parts[pn]["count"] += 1
        unique_parts[pn]["total_quantity"] += item["quantity"] or 0

    # Prepare parts for categorization
    parts_for_ai = [
        {"part_number": p["part_number"], "description": p["description"][:200]}
        for p in unique_parts.values()
    ]

    prompt = f"""Analyze these aerospace/defense parts and categorize each into one of these categories:
- Fasteners (bolts, screws, nuts, rivets, pins)
- Clamps & Retainers (clamps, loops, retainers, brackets)
- Washers & Spacers (washers, spacers, shims)
- Seals & Gaskets (seals, O-rings, gaskets)
- Electrical Components (connectors, wires, terminals)
- Bearings & Bushings (bearings, bushings, sleeves)
- Hardware - Other (miscellaneous hardware)

Return a JSON object mapping each part_number to its category.
Example format: {{"NAS625-19": "Fasteners", "AS21919WDG33": "Clamps & Retainers"}}

Parts to categorize:
{json.dumps(parts_for_ai, indent=2)}

Return ONLY the JSON object, no other text."""

    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.1
    )

    try:
        categorization = json.loads(response.choices[0].message.content)
    except json.JSONDecodeError:
        # Try to extract JSON from response
        content = response.choices[0].message.content
        start = content.find("{")
        end = content.rfind("}") + 1
        categorization = json.loads(content[start:end])

    # Merge categorization with part data
    for pn, data in unique_parts.items():
        data["category"] = categorization.get(pn, "Hardware - Other")

    return unique_parts


def generate_summary_stats(categorized_parts: dict) -> dict:
    """Generate summary statistics by category."""
    category_stats = {}

    for pn, data in categorized_parts.items():
        cat = data["category"]
        if cat not in category_stats:
            category_stats[cat] = {
                "unique_parts": 0,
                "total_appearances": 0,
                "total_quantity": 0,
                "parts": []
            }
        category_stats[cat]["unique_parts"] += 1
        category_stats[cat]["total_appearances"] += data["count"]
        category_stats[cat]["total_quantity"] += data["total_quantity"]
        category_stats[cat]["parts"].append(data)

    # Sort parts within each category by appearance count
    for cat in category_stats:
        category_stats[cat]["parts"].sort(key=lambda x: x["count"], reverse=True)

    return category_stats


def generate_docx_report(category_stats: dict, total_items: int, total_rfqs: int, output_path: str):
    """Generate a Word document report using python-docx."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    # Title
    title = doc.add_heading('RFQ Parts Analysis Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Date
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
    date_run.font.size = Pt(10)
    date_run.font.color.rgb = RGBColor(102, 102, 102)

    # Executive Summary
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(
        f"This report analyzes {total_items} line items across {total_rfqs} RFQ(s). "
        f"Parts have been categorized using AI to identify procurement patterns and frequently requested items."
    )

    # Sort categories by total appearances
    sorted_categories = sorted(
        category_stats.items(),
        key=lambda x: x[1]["total_appearances"],
        reverse=True
    )

    # Category Summary Table
    doc.add_heading('Category Summary', level=1)
    doc.add_paragraph("Parts grouped by category, sorted by frequency of appearance:")

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'

    # Header row
    header_cells = table.rows[0].cells
    headers = ['Category', 'Unique Parts', 'Appearances', 'Total Qty']
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        # Add shading to header
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1F4E79"/>')
        header_cells[i]._tc.get_or_add_tcPr().append(shading)
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    # Data rows
    for cat, stats in sorted_categories:
        row_cells = table.add_row().cells
        row_cells[0].text = cat
        row_cells[1].text = str(stats['unique_parts'])
        row_cells[2].text = str(stats['total_appearances'])
        row_cells[3].text = f"{stats['total_quantity']:,}"

    # Top 15 Most Requested Parts
    doc.add_heading('Top 15 Most Requested Parts', level=1)
    doc.add_paragraph("Individual parts that appear most frequently across RFQs:")

    # Gather all parts and sort by count
    all_parts = []
    for cat, stats in category_stats.items():
        for part in stats["parts"]:
            all_parts.append({**part, "category": cat})
    top_parts = sorted(all_parts, key=lambda x: x["count"], reverse=True)[:15]

    table2 = doc.add_table(rows=1, cols=5)
    table2.style = 'Table Grid'

    # Header row
    header_cells2 = table2.rows[0].cells
    headers2 = ['Part Number', 'Category', 'Appearances', 'Total Qty', 'Description']
    for i, header in enumerate(headers2):
        header_cells2[i].text = header
        header_cells2[i].paragraphs[0].runs[0].bold = True
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1F4E79"/>')
        header_cells2[i]._tc.get_or_add_tcPr().append(shading)
        header_cells2[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    # Data rows
    for part in top_parts:
        row_cells = table2.add_row().cells
        row_cells[0].text = part['part_number']
        row_cells[1].text = part['category']
        row_cells[2].text = str(part['count'])
        row_cells[3].text = f"{part['total_quantity']:,}"
        desc = part.get('description') or ''
        row_cells[4].text = desc[:50] + ('...' if len(desc) > 50 else '')

    # Key Insights
    doc.add_heading('Key Insights', level=1)
    if sorted_categories:
        top_cat, top_stats = sorted_categories[0]
        doc.add_paragraph(
            f"The most common category is \"{top_cat}\" with {top_stats['total_appearances']} appearances across RFQs. "
            f"Consider establishing supplier relationships for high-volume categories to improve pricing and lead times."
        )

    # Save document
    doc.save(output_path)
    print("Report generated successfully")


def main():
    output_path = f"outputs/rfq_parts_analysis_{datetime.now().strftime('%Y-%m-%d')}.docx"
    if len(sys.argv) > 1 and sys.argv[1] == "--output" and len(sys.argv) > 2:
        output_path = sys.argv[2]

    # Ensure output directory exists
    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    print("Connecting to database...")
    engine = get_db_engine()

    print("Fetching line items...")
    line_items = fetch_all_line_items(engine)

    if not line_items:
        print("No line items found in database.")
        sys.exit(0)

    # Get unique RFQ count
    unique_rfqs = len(set(item["rfq_number"] for item in line_items))

    print(f"Found {len(line_items)} line items across {unique_rfqs} RFQ(s)")

    print("Categorizing parts with AI...")
    categorized_parts = categorize_parts_with_ai(line_items)

    print("Generating summary statistics...")
    category_stats = generate_summary_stats(categorized_parts)

    # Print summary to console
    print("\n=== Category Summary ===")
    for cat, stats in sorted(category_stats.items(), key=lambda x: x[1]["total_appearances"], reverse=True):
        print(f"  {cat}: {stats['unique_parts']} unique parts, {stats['total_appearances']} appearances")

    print(f"\nGenerating report: {output_path}")
    generate_docx_report(category_stats, len(line_items), unique_rfqs, output_path)

    print(f"\nDone! Report saved to: {output_path}")


if __name__ == "__main__":
    main()