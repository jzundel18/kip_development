#!/usr/bin/env python3
"""
find_subcontract_opportunities.py

Pulls past contract award winners from SAM.gov and matches them against
your technology areas to surface the best subcontracting opportunities.

How it works:
  1. Fetches award notices from SAM.gov within a configurable lookback window.
  2. Loads your technology area profiles from the database.
  3. Scores every awardee against every technology area using the AI matrix scorer.
  4. Filters to scores above a minimum threshold and ranks by fit.
  5. Prints a ranked report and optionally saves results to a JSON file.

Usage:
    python find_subcontract_opportunities.py [options]

Options:
    --days-back N       How many days of awards to look back (default: 90)
    --max-results N     Maximum matches to show per technology area (default: 10)
    --min-score N       Minimum match score 0-100 to include (default: 65)
    --max-awards N      Maximum awards to pull from SAM.gov (default: 500)
    --output FILE       Save results to a JSON file (default: subcontract_report.json)
    --no-file           Skip saving to file, only print to console
    --naics PREFIX      Filter awards to a NAICS prefix (e.g. 5417 for R&D)
    --verbose           Show detailed scoring breakdown per match

Environment variables (required):
    SUPABASE_DB_URL     PostgreSQL connection string
    SAM_KEYS            Comma-separated SAM.gov API keys
    OPENAI_API_KEY      OpenAI API key for AI scoring

Environment variables (optional):
    DAYS_BACK           Override --days-back
    MAX_RESULTS         Override --max-results
    MIN_SCORE           Override --min-score
    MAX_AWARDS          Override --max-awards
    OUTPUT_FILE         Override --output
"""

from __future__ import annotations

import argparse
import json
import os
import sys
import time
from datetime import datetime, timezone, timedelta
from pathlib import Path
from typing import Any

import requests
import sqlalchemy as sa
from sqlalchemy import text, create_engine

# ── Load .env if available ────────────────────────────────────────────────────
try:
    from dotenv import load_dotenv
    env_file = Path(__file__).parent / ".env"
    if env_file.exists():
        load_dotenv(env_file)
        print(f"✓ Loaded environment from {env_file}")
except ImportError:
    pass

# ── Project path setup ────────────────────────────────────────────────────────
PROJECT_ROOT = Path(__file__).parent
sys.path.insert(0, str(PROJECT_ROOT))
sys.path.insert(0, str(PROJECT_ROOT / "scripts"))

from scoring import AIMatrixScorer  # noqa: E402

# ── Terminal colors ───────────────────────────────────────────────────────────
class C:
    HEADER  = "\033[95m"
    BLUE    = "\033[94m"
    CYAN    = "\033[96m"
    GREEN   = "\033[92m"
    YELLOW  = "\033[93m"
    RED     = "\033[91m"
    BOLD    = "\033[1m"
    END     = "\033[0m"

def h(msg):  print(f"\n{C.HEADER}{C.BOLD}{'─'*68}\n{msg}\n{'─'*68}{C.END}\n")
def ok(msg): print(f"{C.GREEN}✓  {msg}{C.END}")
def warn(msg): print(f"{C.YELLOW}⚠  {msg}{C.END}")
def err(msg): print(f"{C.RED}✗  {msg}{C.END}")
def info(msg): print(f"{C.CYAN}ℹ  {msg}{C.END}")


# ── Config defaults ───────────────────────────────────────────────────────────
DEFAULTS = {
    "days_back":   int(os.getenv("DAYS_BACK",   "7")),
    "max_results": int(os.getenv("MAX_RESULTS", "5")),
    "min_score":   float(os.getenv("MIN_SCORE", "65")),
    "max_awards":  int(os.getenv("MAX_AWARDS",  "500")),
    "output_file": os.getenv("OUTPUT_FILE", "subcontract_report.docx"),
    "naics_filter": os.getenv("NAICS_FILTER", ""),
}

SAM_API_BASE    = "https://api.sam.gov/opportunities/v2/search"
PAGE_SIZE       = 100          # SAM.gov max per page
REQUEST_TIMEOUT = 30
SCORE_BATCH     = 5            # Awardees scored per AI call (smaller = more reliable)
SAM_MAX_WINDOW  = 90           # SAM.gov rejects date ranges over ~90 days in one request

# ── Defense-only filtering ────────────────────────────────────────────────────
# NAICS prefixes that are clearly health / medical / social services — always excluded
HEALTH_NAICS_PREFIXES = (
    "621",   # Ambulatory Health Care
    "622",   # Hospitals
    "623",   # Nursing and Residential Care
    "624",   # Social Assistance
    "3254",  # Pharmaceutical and Medicine Manufacturing
    "3391",  # Medical Equipment and Supplies Manufacturing
    "5242",  # Health and Medical Insurance
)

# Keywords in agency names or titles that signal non-defense health work
HEALTH_KEYWORDS = {
    "health and human services", "hhs", "centers for medicare",
    "centers for disease control", "cdc", "national institutes of health",
    "nih", "food and drug administration", "fda", "indian health service",
    "substance abuse", "samhsa", "cms.gov", "medicaid", "medicare",
    "hospital", "nursing", "clinical trial", "pharmaceutical",
    "public health", "epidemiol", "vaccine", "mental health",
}


def _is_defense_related(award: dict[str, Any]) -> bool:
    """
    Return True if the award is plausibly defense / national-security related.
    Rejects anything that looks like health, medical, or social services.
    """
    naics = award.get("naics_code", "")
    for prefix in HEALTH_NAICS_PREFIXES:
        if naics.startswith(prefix):
            return False

    # Check agency + title for health keywords
    haystack = (
        (award.get("agency") or "") + " " +
        (award.get("title") or "") + " " +
        (award.get("description") or "")[:500]
    ).lower()
    for kw in HEALTH_KEYWORDS:
        if kw in haystack:
            return False

    return True


# ─────────────────────────────────────────────────────────────────────────────
# SAM.GOV AWARD FETCHING
# ─────────────────────────────────────────────────────────────────────────────

def _rotate_keys(keys: list[str]) -> list[str]:
    """Return keys starting from a random index for load distribution."""
    if not keys:
        return keys
    import random
    idx = random.randint(0, len(keys) - 1)
    return keys[idx:] + keys[:idx]


def _parse_award_record(rec: dict[str, Any]) -> dict[str, Any] | None:
    """
    Extract fields we care about from a SAM.gov award notice record.
    Returns None if we can't get a meaningful awardee name.
    """
    def _get(*keys, default=""):
        for k in keys:
            v = rec.get(k)
            if v not in (None, "", [], {}):
                if isinstance(v, (str, int, float)):
                    return str(v).strip()
                if isinstance(v, dict):
                    for sub in ("name", "text", "value", "code"):
                        if v.get(sub) not in (None, "", []):
                            return str(v[sub]).strip()
        return default

    # ── Awardee details ───────────────────────────────────────────────────────
    awardee_name  = ""
    awardee_city  = ""
    awardee_state = ""
    award_amount  = ""
    award_date    = ""

    award_block = rec.get("award") or {}
    if isinstance(award_block, dict):
        awardee_block = award_block.get("awardee") or {}
        if isinstance(awardee_block, dict):
            awardee_name  = str(awardee_block.get("name", "") or "").strip()
            loc = awardee_block.get("location") or {}
            if isinstance(loc, dict):
                awardee_city  = str(loc.get("city", {}).get("name", "") or "").strip()
                awardee_state = str(loc.get("state", {}).get("code", "") or "").strip()
        award_amount = str(award_block.get("amount", "") or "").strip()
        award_date   = str(award_block.get("date", "") or "").strip()

    # Fallback name searches
    if not awardee_name:
        awardee_name = _get("awardeeName", "awardee_name", "company_name")

    if not awardee_name:
        return None   # Can't match without a company name

    # ── Notice fields ─────────────────────────────────────────────────────────
    notice_id       = _get("noticeId", "id")
    title           = _get("title")
    naics_code      = _get("naicsCode", "naics")
    set_aside_code  = _get("typeOfSetAside", "setAsideCode")
    agency          = _get("fullParentPathName", "organizationHierarchy",
                            "agencyName", "department")
    description     = ""

    # Try to get description text
    for key in ("description", "fullParent"):
        block = rec.get(key)
        if isinstance(block, str) and len(block) > 50 and not block.startswith("http"):
            description = block[:3000]
            break
        if isinstance(block, dict):
            d = block.get("description", "")
            if isinstance(d, str) and len(d) > 50:
                description = d[:3000]
                break

    # ── Place of performance ──────────────────────────────────────────────────
    pop = rec.get("placeOfPerformance") or {}
    pop_city  = ""
    pop_state = ""
    if isinstance(pop, dict):
        pop_city  = str(pop.get("city",  {}).get("name",  "") or "").strip()
        pop_state = str(pop.get("state", {}).get("code",  "") or "").strip()

    # ── Public URL ────────────────────────────────────────────────────────────
    link = f"https://sam.gov/opp/{notice_id}/view" if notice_id else "https://sam.gov/"
    for lnk in (rec.get("links") or []):
        if isinstance(lnk, dict) and lnk.get("href"):
            href = str(lnk["href"])
            if "api.sam.gov" not in href:
                link = href
                break

    return {
        "notice_id":      notice_id,
        "title":          title,
        "description":    description,
        "naics_code":     naics_code,
        "set_aside_code": set_aside_code,
        "agency":         agency,
        "awardee_name":   awardee_name,
        "awardee_city":   awardee_city,
        "awardee_state":  awardee_state,
        "award_amount":   award_amount,
        "award_date":     award_date,
        "pop_city":       pop_city,
        "pop_state":      pop_state,
        "link":           link,
    }


def _fetch_awards_window(
    sam_keys: list[str],
    posted_from: str,
    posted_to: str,
    max_awards: int,
    naics_filter: str,
    keys: list[str],
    key_idx_ref: list[int],
    exhausted: set,
) -> tuple[list[dict[str, Any]], int]:
    """
    Fetch awards within a single date window (must be ≤ SAM_MAX_WINDOW days).
    Returns (awards, total_api_records_seen).
    key_idx_ref is a mutable 1-element list so the caller can track rotation state.
    """
    awards     = []
    offset     = 0
    total_seen = 0

    while len(awards) < max_awards:
        if len(exhausted) >= len(keys):
            warn("  All SAM.gov API keys rate-limited or exhausted.")
            break

        key = keys[key_idx_ref[0] % len(keys)]
        if key in exhausted:
            key_idx_ref[0] += 1
            continue

        limit  = min(PAGE_SIZE, max_awards - len(awards))
        params: dict[str, Any] = {
            "api_key":    key,
            "ptype":      "a",
            "postedFrom": posted_from,
            "postedTo":   posted_to,
            "limit":      limit,
            "offset":     offset,
        }
        if naics_filter:
            params["naicsCode"] = naics_filter

        try:
            resp = requests.get(SAM_API_BASE, params=params, timeout=REQUEST_TIMEOUT)
        except requests.exceptions.Timeout:
            warn(f"  Timeout at offset {offset}. Trying next key…")
            exhausted.add(key)
            key_idx_ref[0] += 1
            continue
        except requests.exceptions.RequestException as exc:
            warn(f"  Network error: {exc}")
            break

        if resp.status_code == 429:
            warn(f"  Key {key[:8]}… rate-limited.")
            exhausted.add(key)
            key_idx_ref[0] += 1
            continue

        if resp.status_code == 401:
            warn(f"  Key {key[:8]}… auth failed.")
            exhausted.add(key)
            key_idx_ref[0] += 1
            continue

        if resp.status_code == 400:
            # Surface the actual SAM.gov error message
            try:
                msg = resp.json().get("errorMessage", resp.text[:200])
            except Exception:
                msg = resp.text[:200]
            warn(f"  SAM.gov 400 for window {posted_from}→{posted_to}: {msg}")
            break

        if resp.status_code != 200:
            warn(f"  HTTP {resp.status_code} — skipping window.")
            break

        try:
            data = resp.json()
        except ValueError:
            warn("  Could not parse JSON response.")
            break

        records = data.get("opportunitiesData") or []
        if not records:
            break

        total_seen += len(records)
        batch_added = 0
        seen_ids = {a["notice_id"] for a in awards}

        for rec in records:
            parsed = _parse_award_record(rec)
            if parsed is None:
                continue
            if naics_filter and not parsed["naics_code"].startswith(naics_filter):
                continue
            if parsed["notice_id"] in seen_ids:
                continue
            awards.append(parsed)
            seen_ids.add(parsed["notice_id"])
            batch_added += 1

        print(f"    {posted_from}–{posted_to}  offset={offset:4d}  "
              f"got {len(records):3d}  kept {batch_added:3d}  "
              f"(total so far: {len(awards)})")

        if len(records) < limit:
            break

        offset += PAGE_SIZE
        time.sleep(0.25)

    return awards, total_seen


def fetch_awards(
    sam_keys: list[str],
    days_back: int,
    max_awards: int,
    naics_filter: str = "",
) -> list[dict[str, Any]]:
    """
    Pull award notices from SAM.gov for the past `days_back` days.
    SAM.gov rejects date ranges > ~90 days, so we chunk automatically.
    """
    h("STEP 1 — Fetching Award Notices from SAM.gov")

    end_dt   = datetime.now(timezone.utc)
    start_dt = end_dt - timedelta(days=days_back)

    info(f"Total range  : {start_dt.strftime('%m/%d/%Y')}  →  {end_dt.strftime('%m/%d/%Y')}  ({days_back} days)")
    info(f"Window size  : {SAM_MAX_WINDOW} days per request (SAM.gov limit)")
    info(f"Max awards   : {max_awards}")
    if naics_filter:
        info(f"NAICS filter : {naics_filter}*")

    # Build non-overlapping date windows from newest → oldest
    windows: list[tuple[str, str]] = []
    window_end = end_dt
    while window_end > start_dt:
        window_start = max(window_end - timedelta(days=SAM_MAX_WINDOW), start_dt)
        windows.append((
            window_start.strftime("%m/%d/%Y"),
            window_end.strftime("%m/%d/%Y"),
        ))
        window_end = window_start
    info(f"Windows      : {len(windows)}  ({SAM_MAX_WINDOW}-day chunks)\n")

    keys        = _rotate_keys(sam_keys)
    key_idx_ref = [0]
    exhausted: set[str] = set()
    all_awards: list[dict[str, Any]] = []
    all_seen_ids: set[str] = set()
    total_api_seen = 0

    for win_from, win_to in windows:
        if len(all_awards) >= max_awards:
            info("Reached max_awards cap — stopping early.")
            break
        if len(exhausted) >= len(keys):
            warn("All API keys exhausted — stopping early.")
            break

        remaining = max_awards - len(all_awards)
        window_awards, seen = _fetch_awards_window(
            sam_keys=sam_keys,
            posted_from=win_from,
            posted_to=win_to,
            max_awards=remaining,
            naics_filter=naics_filter,
            keys=keys,
            key_idx_ref=key_idx_ref,
            exhausted=exhausted,
        )
        total_api_seen += seen

        # Global dedup across windows
        for aw in window_awards:
            if aw["notice_id"] not in all_seen_ids:
                all_awards.append(aw)
                all_seen_ids.add(aw["notice_id"])

    ok(f"Fetched {len(all_awards)} unique award records "
       f"(from {total_api_seen} total returned by SAM.gov across {len(windows)} window(s))")
    return all_awards


# ─────────────────────────────────────────────────────────────────────────────
# DATABASE — TECHNOLOGY AREAS
# ─────────────────────────────────────────────────────────────────────────────

def load_technology_areas(db_url: str) -> list[dict[str, Any]]:
    """Load all technology areas with descriptions from the database."""
    h("STEP 2 — Loading Technology Profiles from Database")

    engine = create_engine(db_url, pool_pre_ping=True)
    try:
        with engine.connect() as conn:
            rows = conn.execute(text("""
                SELECT id,
                       technology_name,
                       description,
                       emails
                FROM technology_areas
                WHERE description IS NOT NULL
                  AND description != ''
                ORDER BY technology_name
            """)).fetchall()
    except Exception as exc:
        err(f"Could not read technology_areas: {exc}")
        sys.exit(1)

    # Only include Dynamic Camo and Dynamic Infrared (DIR) programs
    ALLOWED_TECH_KEYWORDS = ("dynamic camo", "dynamic infrared", "dir")

    def _is_allowed(name: str) -> bool:
        n = name.lower()
        # Match "dynamic camo", "dynamic infrared", or "DIR" as a standalone token/program
        if "dynamic camo" in n or "dynamic infrared" in n:
            return True
        # "dir" as a standalone word or in parentheses like "(DIR)"
        import re
        if re.search(r"\bdir\b", n):
            return True
        return False

    areas = [
        {
            "id":          str(r[0]),
            "name":        str(r[1] or "").strip(),
            "description": str(r[2] or "").strip(),
            "emails":      str(r[3] or "").strip(),
        }
        for r in rows
        if (r[1] or "").strip()
           and (r[2] or "").strip()
           and _is_allowed(str(r[1] or ""))
    ]

    if not areas:
        err("No matching technology areas found. Expected 'Dynamic Camo' "
            "and/or 'Dynamic Infrared (DIR)' in the technology_areas table.")
        sys.exit(1)

    ok(f"Loaded {len(areas)} technology area(s):")
    for a in areas:
        print(f"   • {a['name']}")
    return areas


# ─────────────────────────────────────────────────────────────────────────────
# SCORING
# ─────────────────────────────────────────────────────────────────────────────

def _build_scoring_profile(tech_area: dict[str, Any]) -> dict[str, str]:
    """
    Shape a technology area dict into the profile dict AIMatrixScorer expects.
    We borrow the scoring format used for solicitations, treating the award
    title + description as the 'solicitation' and the tech area as the 'company'.
    """
    return {
        "company_name": tech_area["name"],
        "description":  tech_area["description"],
        "city":         "",
        "state":        "",
    }


def score_awards_for_technology(
    awards: list[dict[str, Any]],
    tech_area: dict[str, Any],
    openai_key: str,
    min_score: float,
    max_results: int,
    verbose: bool = False,
) -> list[dict[str, Any]]:
    """
    Score every award awardee against a single technology area profile.
    Returns matches above `min_score`, sorted descending, capped at `max_results`.
    """
    import pandas as pd  # local import — only needed here

    if not awards:
        return []

    # Filter to defense-related awards only
    defense_awards = [aw for aw in awards if _is_defense_related(aw)]
    if not defense_awards:
        warn(f"  No defense-related awards found for scoring.")
        return []
    if len(defense_awards) < len(awards):
        info(f"  Filtered {len(awards)} → {len(defense_awards)} defense-related awards "
             f"(excluded {len(awards) - len(defense_awards)} health/non-defense)")

    scorer  = AIMatrixScorer()
    profile = _build_scoring_profile(tech_area)

    # Build a lightweight DataFrame so score_batch can consume it
    rows = []
    for aw in defense_awards:
        # For scoring: title = "[AwardeeCompany] — [Solicitation Title]"
        # description = agency + work description
        combined_title = f"{aw['awardee_name']} — {aw['title']}"
        combined_desc  = (
            f"Awardee: {aw['awardee_name']}. "
            f"Agency: {aw['agency']}. "
            f"NAICS: {aw['naics_code']}. "
            f"{aw['description']}"
        )[:2000]
        rows.append({
            "notice_id":      aw["notice_id"],
            "title":          combined_title,
            "description":    combined_desc,
            "naics_code":     aw["naics_code"],
            "set_aside_code": aw["set_aside_code"],
            "pop_city":       aw["pop_city"],
            "pop_state":      aw["pop_state"],
        })

    df = pd.DataFrame(rows)
    all_scores: dict[str, dict] = {}

    # Score in small batches to stay reliable
    for batch_start in range(0, len(rows), SCORE_BATCH):
        batch = rows[batch_start : batch_start + SCORE_BATCH]
        try:
            batch_scores = scorer.score_batch(batch, company_profile=profile,
                                              api_key=openai_key, model="gpt-4o-mini")
            all_scores.update(batch_scores)
        except Exception as exc:
            warn(f"Scoring batch {batch_start // SCORE_BATCH + 1} failed: {exc}")
        time.sleep(0.2)

    # Map scores back to award records
    results = []
    for aw in defense_awards:
        nid   = aw["notice_id"]
        entry = all_scores.get(nid, {})
        score = float(entry.get("score", 0.0))
        if score < min_score:
            continue
        results.append({
            **aw,
            "score":     round(score, 1),
            "breakdown": entry.get("breakdown", []),
            "tech_area": tech_area["name"],
        })

    results.sort(key=lambda x: x["score"], reverse=True)
    top_results = results[:max_results]

    # Generate AI rationale paragraphs for the top matches
    if top_results:
        _generate_rationales(top_results, tech_area, openai_key)

    return top_results


def _generate_rationales(
    matches: list[dict[str, Any]],
    tech_area: dict[str, Any],
    openai_key: str,
):
    """
    For each match, call the AI to produce a short paragraph explaining
    why the technology area is a strong subcontracting fit for that award.
    Adds a 'rationale' key to each match dict in-place.
    """
    import openai

    client = openai.OpenAI(api_key=openai_key)

    for m in matches:
        prompt = (
            f"You are a defense contracting analyst. Write ONE concise paragraph (3-5 sentences) "
            f"explaining why a company specializing in \"{tech_area['name']}\" "
            f"({tech_area['description'][:300]}) "
            f"would be an excellent subcontracting partner for the prime contractor "
            f"\"{m['awardee_name']}\" on this defense contract:\n\n"
            f"Title: {m['title']}\n"
            f"Agency: {m['agency']}\n"
            f"NAICS: {m['naics_code']}\n"
            f"Description: {(m.get('description') or '')[:500]}\n\n"
            f"Focus on specific technical synergies, capability gaps the subcontractor "
            f"could fill, and why this pairing makes strategic sense for winning and "
            f"executing the work. Be specific and actionable, not generic."
        )

        try:
            resp = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=300,
                temperature=0.7,
            )
            m["rationale"] = resp.choices[0].message.content.strip()
        except Exception as exc:
            warn(f"  Rationale generation failed for {m['awardee_name']}: {exc}")
            m["rationale"] = ""
        time.sleep(0.2)


# ─────────────────────────────────────────────────────────────────────────────
# REPORT RENDERING
# ─────────────────────────────────────────────────────────────────────────────

def _score_bar(score: float, width: int = 30) -> str:
    filled = int(round(score / 100 * width))
    bar    = "█" * filled + "░" * (width - filled)
    color  = C.GREEN if score >= 80 else C.CYAN if score >= 65 else C.YELLOW
    return f"{color}{bar}{C.END}  {C.BOLD}{score:5.1f}{C.END}"


def print_report(
    results_by_tech: dict[str, list[dict[str, Any]]],
    config: dict[str, Any],
    verbose: bool = False,
):
    h("SUBCONTRACTING OPPORTUNITY REPORT")

    total_matches = sum(len(v) for v in results_by_tech.values())
    info(f"Configuration: days_back={config['days_back']}  "
         f"min_score={config['min_score']}  max_results={config['max_results']}")
    print()

    if total_matches == 0:
        warn("No matches found above the minimum score threshold.")
        info("Try lowering --min-score or increasing --days-back.")
        return

    ok(f"Found {total_matches} subcontracting opportunities across "
       f"{sum(1 for v in results_by_tech.values() if v)} technology area(s).\n")

    for tech_name, matches in results_by_tech.items():
        if not matches:
            continue

        print(f"\n{C.BOLD}{C.HEADER}━━━  {tech_name.upper()}  ({len(matches)} matches){C.END}")
        print()

        for rank, m in enumerate(matches, 1):
            amount_str = f"  ${float(m['award_amount']):,.0f}" if m["award_amount"] else ""
            date_str   = f"  Awarded: {m['award_date']}" if m["award_date"] else ""
            loc_parts  = [p for p in [m["awardee_city"], m["awardee_state"]] if p]
            loc_str    = f"  📍 {', '.join(loc_parts)}" if loc_parts else ""

            print(f"  {C.BOLD}#{rank}  {m['awardee_name']}{C.END}")
            print(f"      Score:  {_score_bar(m['score'])}")
            print(f"      Work:   {m['title'][:90]}")
            if m["agency"]:
                print(f"      Agency: {m['agency'][:80]}")
            print(f"      NAICS:  {m['naics_code']}{amount_str}{date_str}{loc_str}")
            print(f"      Link:   {m['link']}")

            if m.get("rationale"):
                # Word-wrap rationale to ~70 chars with indent
                import textwrap
                wrapped = textwrap.fill(m["rationale"], width=70,
                                        initial_indent="      ",
                                        subsequent_indent="      ")
                print(f"\n{C.CYAN}{wrapped}{C.END}\n")

            if verbose and m.get("breakdown"):
                print(f"      {C.CYAN}Score breakdown:{C.END}")
                for comp in m["breakdown"]:
                    bar_mini = "█" * comp["score"] + "░" * (10 - comp["score"])
                    print(f"        {bar_mini}  {comp['score']:2d}/10  "
                          f"{comp['label'][:40]:<40}  {comp.get('reasoning','')[:50]}")
            print()


# ─────────────────────────────────────────────────────────────────────────────
# DOCX REPORT GENERATION
# ─────────────────────────────────────────────────────────────────────────────

def _score_label(score: float) -> str:
    if score >= 85: return "Excellent Fit"
    if score >= 75: return "Strong Fit"
    if score >= 65: return "Good Fit"
    return "Moderate Fit"


def _score_rgb(score: float):
    from docx.shared import RGBColor
    if score >= 85: return RGBColor(0x1D, 0x6F, 0x42)
    if score >= 75: return RGBColor(0x2E, 0x75, 0xB6)
    if score >= 65: return RGBColor(0xC0, 0x70, 0x00)
    return RGBColor(0x7F, 0x7F, 0x7F)


def generate_docx_report(
    results_by_tech: dict[str, list[dict[str, Any]]],
    config: dict[str, Any],
    output_path: str,
):
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    generated_at  = datetime.now().strftime("%B %d, %Y  %H:%M")
    date_range    = (
        f"Past {config['days_back']} days  |  "
        f"Min score: {config['min_score']}  |  "
        f"Max results per area: {config['max_results']}"
    )
    total_matches = sum(len(v) for v in results_by_tech.values())

    # ── Cover header ──────────────────────────────────────────────────────────
    title = doc.add_heading('KIP Subcontract Opportunity Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for txt, bold in [
        (f"Generated: {generated_at}", False),
        (date_range, False),
        (f"Total matches found: {total_matches}", True),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(txt)
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
        run.bold = bold

    # ── One section per technology area ───────────────────────────────────────
    for tech_name, matches in results_by_tech.items():
        if not matches:
            continue

        doc.add_heading(tech_name, level=1)

        opp_word = "opportunity" if len(matches) == 1 else "opportunities"
        sub = doc.add_paragraph()
        run = sub.add_run(f"{len(matches)} subcontracting {opp_word} identified")
        run.italic = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

        for rank, m in enumerate(matches, 1):
            score   = m["score"]
            label   = _score_label(score)
            color   = _score_rgb(score)

            try:
                amount_display = f"${float(m['award_amount']):,.0f}" if m.get("award_amount") else "Not disclosed"
            except (ValueError, TypeError):
                amount_display = m.get("award_amount") or "Not disclosed"

            award_date = m.get("award_date") or "Not specified"
            loc_parts  = [p for p in [m.get("awardee_city", ""), m.get("awardee_state", "")] if p]
            location   = ", ".join(loc_parts) if loc_parts else "Location not specified"
            agency     = m.get("agency") or "Not specified"
            naics      = m.get("naics_code") or "N/A"
            link       = m.get("link") or "https://sam.gov/"
            title_text = (m.get("title") or "Untitled")[:120]
            awardee    = m.get("awardee_name") or "Unknown"

            # ── Awardee heading ───────────────────────────────────────────────
            doc.add_heading(f"#{rank}  {awardee}", level=2)

            # ── Score line ────────────────────────────────────────────────────
            score_para = doc.add_paragraph()
            r1 = score_para.add_run("Match Score: ")
            r1.bold = True
            r2 = score_para.add_run(f"{score:.1f} / 100")
            r2.bold = True
            r2.font.color.rgb = color
            r3 = score_para.add_run(f"  \u2014  {label}")
            r3.italic = True
            r3.font.color.rgb = color

            # ── Details table ─────────────────────────────────────────────────
            detail_rows = [
                ("Contract / Work",  title_text),
                ("Awarding Agency",  agency[:100]),
                ("Award Amount",     amount_display),
                ("Award Date",       award_date),
                ("Company Location", location),
                ("NAICS Code",       naics),
                ("SAM.gov Link",     link),
            ]

            table = doc.add_table(rows=len(detail_rows), cols=2)
            table.style = 'Table Grid'

            for i, (lbl, val) in enumerate(detail_rows):
                label_cell = table.rows[i].cells[0]
                value_cell = table.rows[i].cells[1]

                label_cell.text = lbl
                label_cell.paragraphs[0].runs[0].bold = True
                label_cell.paragraphs[0].runs[0].font.size = Pt(10)
                # Shade the label column
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="EBF3FB"/>')
                label_cell._tc.get_or_add_tcPr().append(shading)

                value_cell.text = val
                value_cell.paragraphs[0].runs[0].font.size = Pt(10)

            # ── Rationale paragraph ────────────────────────────────────────
            if m.get("rationale"):
                rat_heading = doc.add_paragraph()
                rh_run = rat_heading.add_run("Why This Is a Strong Subcontracting Fit")
                rh_run.bold = True
                rh_run.font.size = Pt(11)

                rat_para = doc.add_paragraph()
                rat_run = rat_para.add_run(m["rationale"])
                rat_run.font.size = Pt(10)
                rat_run.italic = True

            doc.add_paragraph()  # spacer

            # ── Score breakdown ───────────────────────────────────────────────
            if m.get("breakdown"):
                bd_heading = doc.add_paragraph()
                bd_run = bd_heading.add_run("Score Breakdown")
                bd_run.bold = True

                bd_table = doc.add_table(rows=1, cols=3)
                bd_table.style = 'Table Grid'

                # Header row
                headers = ['Scoring Category', 'Score', 'AI Reasoning']
                for j, hdr in enumerate(headers):
                    cell = bd_table.rows[0].cells[j]
                    cell.text = hdr
                    cell.paragraphs[0].runs[0].bold = True
                    cell.paragraphs[0].runs[0].font.size = Pt(9)
                    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F2F2"/>')
                    cell._tc.get_or_add_tcPr().append(shading)

                for comp in m["breakdown"]:
                    cs   = comp.get("score", 0)
                    cc   = _score_rgb(cs * 10)  # scale 0-10 to 0-100 range
                    row  = bd_table.add_row()
                    row.cells[0].text = comp.get("label", "")
                    row.cells[0].paragraphs[0].runs[0].font.size = Pt(9)

                    score_cell = row.cells[1]
                    score_cell.text = f"{cs}/10"
                    score_cell.paragraphs[0].runs[0].bold = True
                    score_cell.paragraphs[0].runs[0].font.size = Pt(9)
                    score_cell.paragraphs[0].runs[0].font.color.rgb = cc

                    row.cells[2].text = comp.get("reasoning", "")
                    row.cells[2].paragraphs[0].runs[0].font.size = Pt(9)

                doc.add_paragraph()  # spacer

    # ── Save ──────────────────────────────────────────────────────────────────
    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
    doc.save(output_path)
    ok(f"Word document saved to: {output_path}")


# ─────────────────────────────────────────────────────────────────────────────
# CLI + MAIN
# ─────────────────────────────────────────────────────────────────────────────

def parse_args() -> argparse.Namespace:
    p = argparse.ArgumentParser(
        description="Find subcontracting opportunities from SAM.gov award winners.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=__doc__,
    )
    p.add_argument("--days-back",   type=int,   default=DEFAULTS["days_back"],
                   help=f"Days of history to search (default: {DEFAULTS['days_back']})")
    p.add_argument("--max-results", type=int,   default=DEFAULTS["max_results"],
                   help=f"Max matches per technology area (default: {DEFAULTS['max_results']})")
    p.add_argument("--min-score",   type=float, default=DEFAULTS["min_score"],
                   help=f"Minimum match score 0-100 (default: {DEFAULTS['min_score']})")
    p.add_argument("--max-awards",  type=int,   default=DEFAULTS["max_awards"],
                   help=f"Max awards to pull from SAM.gov (default: {DEFAULTS['max_awards']})")
    p.add_argument("--output",      type=str,   default=DEFAULTS["output_file"],
                   help=f"Output .docx file path (default: {DEFAULTS['output_file']})")
    p.add_argument("--no-file",     action="store_true",
                   help="Skip saving .docx report")
    p.add_argument("--naics",       type=str,   default=DEFAULTS["naics_filter"],
                   help="Filter to awards with this NAICS prefix (e.g. '5417')")
    p.add_argument("--verbose",     action="store_true",
                   help="Show per-component score breakdown for each match")
    return p.parse_args()


def main():
    args = parse_args()

    # ── Banner ────────────────────────────────────────────────────────────────
    print(f"\n{C.BOLD}{C.HEADER}")
    print("╔══════════════════════════════════════════════════════════════════╗")
    print("║          KIP  —  SUBCONTRACT OPPORTUNITY FINDER                 ║")
    print("╚══════════════════════════════════════════════════════════════════╝")
    print(f"{C.END}")
    print(f"  Started: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    config = {
        "days_back":    args.days_back,
        "max_results":  args.max_results,
        "min_score":    args.min_score,
        "max_awards":   args.max_awards,
        "naics_filter": args.naics,
    }

    # ── Validate environment ──────────────────────────────────────────────────
    h("Environment Check")

    db_url     = os.getenv("SUPABASE_DB_URL")
    openai_key = os.getenv("OPENAI_API_KEY")
    sam_keys_raw = os.getenv("SAM_KEYS", "")
    sam_keys = [k.strip() for k in sam_keys_raw.replace("\n", ",").split(",") if k.strip()]

    missing = []
    if not db_url:       missing.append("SUPABASE_DB_URL")
    if not openai_key:   missing.append("OPENAI_API_KEY")
    if not sam_keys:     missing.append("SAM_KEYS")

    if missing:
        err(f"Missing required env vars: {', '.join(missing)}")
        sys.exit(1)

    ok(f"SUPABASE_DB_URL  — connected")
    ok(f"OPENAI_API_KEY   — configured")
    ok(f"SAM_KEYS         — {len(sam_keys)} key(s)")

    # ── Step 1: Fetch awards ──────────────────────────────────────────────────
    awards = fetch_awards(
        sam_keys=sam_keys,
        days_back=args.days_back,
        max_awards=args.max_awards,
        naics_filter=args.naics,
    )

    if not awards:
        warn("No award records returned. Check your SAM.gov keys and date range.")
        sys.exit(0)

    # ── Step 2: Load technology areas ─────────────────────────────────────────
    tech_areas = load_technology_areas(db_url)

    # ── Step 3: Score awards against each technology area ─────────────────────
    h("STEP 3 — Scoring Awards Against Technology Profiles")
    info(f"Scoring {len(awards)} awards × {len(tech_areas)} technology area(s)…")
    info(f"This uses gpt-4o-mini in small batches — may take a few minutes.\n")

    results_by_tech: dict[str, list[dict[str, Any]]] = {}

    for idx, tech in enumerate(tech_areas, 1):
        print(f"\n[{idx}/{len(tech_areas)}]  Scoring for:  {C.BOLD}{tech['name']}{C.END}")

        matches = score_awards_for_technology(
            awards=awards,
            tech_area=tech,
            openai_key=openai_key,
            min_score=args.min_score,
            max_results=args.max_results,
            verbose=args.verbose,
        )

        results_by_tech[tech["name"]] = matches
        ok(f"  → {len(matches)} matches above score {args.min_score}")

    # ── Step 4: Print report ──────────────────────────────────────────────────
    print_report(results_by_tech, config, verbose=args.verbose)

    # ── Step 5: Generate .docx report ────────────────────────────────────────
    if not args.no_file:
        h("STEP 5 — Generating Word Document Report")
        # Ensure .docx extension
        out_name = args.output
        if not out_name.lower().endswith(".docx"):
            out_name = out_name.rsplit(".", 1)[0] + ".docx" if "." in out_name else out_name + ".docx"
        out_path = str(Path(out_name).resolve())
        try:
            generate_docx_report(results_by_tech, config, out_path)
        except Exception as exc:
            err(f"Could not generate .docx report: {exc}")
            import traceback
            traceback.print_exc()

    # ── Summary ───────────────────────────────────────────────────────────────
    h("COMPLETE")
    total = sum(len(v) for v in results_by_tech.values())
    ok(f"Total subcontracting opportunities found: {total}")
    print(f"  Finished: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    print()


if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        print(f"\n\n{C.YELLOW}Interrupted by user.{C.END}")
        sys.exit(130)
    except Exception as exc:
        err(f"Fatal error: {exc}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
